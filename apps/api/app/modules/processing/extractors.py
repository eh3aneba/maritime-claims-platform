from __future__ import annotations

import hashlib
import shutil
import subprocess
import tempfile
import time
from dataclasses import dataclass, field
from pathlib import Path

EXTRACTOR_VERSION = "2.1"
MIN_MEANINGFUL_TEXT_CHARS = 40
MIN_MEANINGFUL_PAGE_TEXT_CHARS = 20


@dataclass(frozen=True)
class TextSegmentResult:
    locator_type: str
    locator_value: str
    text: str


@dataclass(frozen=True)
class TextExtractionResult:
    method: str
    segments: list[TextSegmentResult] = field(default_factory=list)
    requires_ocr: bool = False
    warnings: list[str] = field(default_factory=list)

    @property
    def combined_text(self) -> str:
        return "\n\n".join(segment.text for segment in self.segments if segment.text.strip())

    @property
    def text_hash(self) -> str | None:
        text = self.combined_text.strip()
        return hashlib.sha256(text.encode("utf-8")).hexdigest() if text else None


def extract_document_text(
    path: Path,
    *,
    enable_ocr: bool = False,
    ocr_languages: str = "eng+fas",
    ocr_max_pages: int = 20,
    ocr_timeout_seconds: float = 120.0,
) -> TextExtractionResult:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        native = _extract_pdf(path)
        if enable_ocr and native.requires_ocr:
            return _extract_pdf_ocr(
                path,
                native_result=native,
                languages=ocr_languages,
                max_pages=ocr_max_pages,
                timeout_seconds=ocr_timeout_seconds,
            )
        return native
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".xlsx":
        return _extract_xlsx(path)
    if suffix in {".jpg", ".jpeg", ".png"}:
        if enable_ocr:
            return _extract_image_ocr(
                path,
                languages=ocr_languages,
                timeout_seconds=ocr_timeout_seconds,
            )
        return TextExtractionResult(
            method="image_requires_ocr",
            requires_ocr=True,
            warnings=["Image documents require OCR; local OCR is not enabled for this process."],
        )
    raise ValueError(f"No text extractor is available for {suffix or 'this file type'}")


def _tesseract_available() -> bool:
    return shutil.which("tesseract") is not None


def _ocr_one_image(path: Path, *, languages: str, timeout_seconds: float) -> str:
    completed = subprocess.run(
        ["tesseract", str(path), "stdout", "-l", languages],
        check=True,
        capture_output=True,
        text=True,
        timeout=timeout_seconds,
    )
    return completed.stdout.strip()


def _extract_image_ocr(
    path: Path,
    *,
    languages: str,
    timeout_seconds: float,
) -> TextExtractionResult:
    if not _tesseract_available():
        return TextExtractionResult(
            method="tesseract_unavailable",
            requires_ocr=True,
            warnings=["Local OCR is enabled but the Tesseract binary is unavailable."],
        )
    try:
        text = _ocr_one_image(path, languages=languages, timeout_seconds=timeout_seconds)
    except subprocess.TimeoutExpired:
        return TextExtractionResult(
            method="tesseract_timeout",
            requires_ocr=True,
            warnings=["Local OCR timed out before producing reviewable text."],
        )
    except (subprocess.SubprocessError, OSError) as exc:
        return TextExtractionResult(
            method="tesseract_failed",
            requires_ocr=True,
            warnings=[f"Local OCR failed: {type(exc).__name__}."],
        )
    return TextExtractionResult(
        method=f"tesseract:{languages}",
        segments=[TextSegmentResult(locator_type="page", locator_value="1", text=text)],
        requires_ocr=not bool(text.strip()),
        warnings=[] if text.strip() else ["OCR completed but produced no meaningful text."],
    )


def _low_text_page_numbers(result: TextExtractionResult) -> list[int]:
    low: list[int] = []
    for index, segment in enumerate(result.segments, start=1):
        if len(segment.text.strip()) < MIN_MEANINGFUL_PAGE_TEXT_CHARS:
            try:
                low.append(int(segment.locator_value))
            except ValueError:
                low.append(index)
    return low


def _extract_pdf_ocr(
    path: Path,
    *,
    native_result: TextExtractionResult | None = None,
    languages: str,
    max_pages: int,
    timeout_seconds: float,
) -> TextExtractionResult:
    """OCR only low-text PDF pages while preserving useful native page text.

    The configured page cap is a document page ceiling: low-text pages beyond it
    stay unresolved and are surfaced explicitly. A timeout/failure preserves all
    native text plus any OCR pages already completed, so partial evidence never
    masquerades as a complete extraction and useful text is not discarded.
    """

    native = native_result or _extract_pdf(path)
    low_pages = _low_text_page_numbers(native)
    if not low_pages:
        return native

    if not _tesseract_available() or shutil.which("pdftoppm") is None:
        return TextExtractionResult(
            method="pypdf+ocr_unavailable",
            segments=native.segments,
            requires_ocr=True,
            warnings=[
                f"{len(low_pages)} PDF page(s) contain little or no native text.",
                "Selective scanned-PDF OCR requires both Tesseract and pdftoppm.",
            ],
        )

    max_pages = max(1, min(max_pages, 100))
    eligible_pages = [page for page in low_pages if page <= max_pages]
    capped_pages = [page for page in low_pages if page > max_pages]
    deadline = time.monotonic() + max(1.0, timeout_seconds)
    segment_by_page = {
        int(segment.locator_value): segment
        for segment in native.segments
        if segment.locator_type == "page" and segment.locator_value.isdigit()
    }
    warnings: list[str] = []
    completed_pages: list[int] = []
    failed_pages: list[int] = []

    def remaining_seconds() -> float:
        remaining = deadline - time.monotonic()
        if remaining <= 0:
            raise subprocess.TimeoutExpired(cmd="selective-pdf-ocr", timeout=timeout_seconds)
        return remaining

    with tempfile.TemporaryDirectory(prefix="mcri-ocr-") as directory:
        for page_number in eligible_pages:
            prefix = Path(directory) / f"page-{page_number}"
            page_image = prefix.with_suffix(".png")
            try:
                subprocess.run(
                    [
                        "pdftoppm",
                        "-f",
                        str(page_number),
                        "-l",
                        str(page_number),
                        "-singlefile",
                        "-r",
                        "200",
                        "-png",
                        str(path),
                        str(prefix),
                    ],
                    check=True,
                    capture_output=True,
                    timeout=remaining_seconds(),
                )
                ocr_text = _ocr_one_image(
                    page_image,
                    languages=languages,
                    timeout_seconds=remaining_seconds(),
                )
            except subprocess.TimeoutExpired:
                failed_pages.extend(page for page in eligible_pages if page not in completed_pages)
                warnings.append(
                    f"Selective PDF OCR timed out after {len(completed_pages)} page(s); partial extraction was preserved."
                )
                break
            except (subprocess.SubprocessError, OSError) as exc:
                failed_pages.append(page_number)
                warnings.append(
                    f"Selective PDF OCR failed on page {page_number}: {type(exc).__name__}; native/partial text was preserved."
                )
                continue

            if ocr_text.strip():
                segment_by_page[page_number] = TextSegmentResult(
                    locator_type="page",
                    locator_value=str(page_number),
                    text=ocr_text.strip(),
                )
                completed_pages.append(page_number)
            else:
                failed_pages.append(page_number)
                warnings.append(f"OCR page {page_number} produced no meaningful text.")

    if capped_pages:
        warnings.append(
            f"OCR page cap {max_pages} left {len(capped_pages)} low-text page(s) unprocessed: "
            + ", ".join(str(page) for page in capped_pages[:10])
            + ("…" if len(capped_pages) > 10 else "")
            + "."
        )

    merged_segments = [
        segment_by_page.get(index, segment)
        for index, segment in enumerate(native.segments, start=1)
    ]
    unresolved_pages = [
        page
        for page in low_pages
        if len(segment_by_page.get(page, TextSegmentResult("page", str(page), "")).text.strip())
        < MIN_MEANINGFUL_PAGE_TEXT_CHARS
    ]
    if unresolved_pages and not warnings:
        warnings.append(
            f"{len(unresolved_pages)} low-text PDF page(s) still require review/OCR."
        )
    if completed_pages:
        warnings.insert(
            0,
            f"Selective local OCR recovered {len(completed_pages)} low-text PDF page(s) using {languages}.",
        )

    return TextExtractionResult(
        method=f"pypdf+selective-tesseract:{languages}",
        segments=merged_segments,
        requires_ocr=bool(unresolved_pages),
        warnings=warnings,
    )


def _extract_pdf(path: Path) -> TextExtractionResult:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    segments: list[TextSegmentResult] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        segments.append(TextSegmentResult(locator_type="page", locator_value=str(index), text=text))
    low_pages = [
        index
        for index, segment in enumerate(segments, start=1)
        if len(segment.text.strip()) < MIN_MEANINGFUL_PAGE_TEXT_CHARS
    ]
    total_chars = sum(len(segment.text.strip()) for segment in segments)
    requires_ocr = bool(low_pages) or total_chars < MIN_MEANINGFUL_TEXT_CHARS
    warnings: list[str] = []
    if low_pages:
        warnings.append(
            f"PDF contains {len(low_pages)} low-text page(s) that may require OCR."
        )
    elif total_chars < MIN_MEANINGFUL_TEXT_CHARS:
        warnings.append("PDF contains little extractable text and likely requires OCR.")
    return TextExtractionResult(
        method="pypdf", segments=segments, requires_ocr=requires_ocr, warnings=warnings
    )


def _extract_docx(path: Path) -> TextExtractionResult:
    from docx import Document as WordDocument

    document = WordDocument(str(path))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            blocks.append(f"[Table {table_index}]\n" + "\n".join(rows))
    text = "\n\n".join(blocks)
    return TextExtractionResult(
        method="python-docx",
        segments=[TextSegmentResult(locator_type="document", locator_value="body", text=text)],
        requires_ocr=False,
    )


def _extract_xlsx(path: Path) -> TextExtractionResult:
    from openpyxl import load_workbook

    workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
    segments: list[TextSegmentResult] = []
    for worksheet in workbook.worksheets:
        lines: list[str] = []
        for row in worksheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value for value in values):
                lines.append(" | ".join(values))
        segments.append(
            TextSegmentResult(
                locator_type="sheet", locator_value=worksheet.title[:100], text="\n".join(lines)
            )
        )
    workbook.close()
    return TextExtractionResult(method="openpyxl", segments=segments, requires_ocr=False)
