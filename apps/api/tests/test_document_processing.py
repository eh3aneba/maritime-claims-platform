from datetime import date
from io import BytesIO
from pathlib import Path
from uuid import UUID

from docx import Document as WordDocument
from sqlalchemy import select

from app.ai.gateway.base import AIProviderUnavailable, AIRequest
from app.ai.gateway.registry import get_ai_provider
from app.core.security import hash_password
from app.modules.claims.models import Claim
from app.modules.documents.models import Document, DocumentProcessingStatus
from app.modules.documents import service as document_service
from app.modules.organizations.models import Organization
from app.modules.processing.models import (
    DocumentProcessingJob,
    DocumentTextExtraction,
    DocumentTextSegment,
    ProcessingJobStatus,
)
from app.modules.processing.service import claim_next_job, process_job
from app.modules.users.models import User, UserRole
from app.modules.vessels.models import Vessel
from tests.db_harness import TestingSessionLocal, client, reset_database

PASSWORD = "Strong-Processing-Test-2026"


def setup_function() -> None:
    reset_database()


def seed_claim() -> dict[str, str]:
    with TestingSessionLocal() as db:
        alpha = Organization(name="Alpha Marine", slug="alpha")
        beta = Organization(name="Beta Marine", slug="beta")
        db.add_all([alpha, beta])
        db.flush()
        alpha_user = User(
            organization_id=alpha.id,
            email="alpha@example.com",
            full_name="Alpha Handler",
            password_hash=hash_password(PASSWORD),
            role=UserRole.CLAIMS_HANDLER,
            is_active=True,
        )
        beta_user = User(
            organization_id=beta.id,
            email="beta@example.com",
            full_name="Beta Handler",
            password_hash=hash_password(PASSWORD),
            role=UserRole.CLAIMS_HANDLER,
            is_active=True,
        )
        vessel = Vessel(organization_id=alpha.id, name="MT ORION", imo_number="7000101")
        beta_vessel = Vessel(organization_id=beta.id, name="MT BETA", imo_number="7000102")
        db.add_all([alpha_user, beta_user, vessel, beta_vessel])
        db.flush()
        claim = Claim(
            organization_id=alpha.id,
            vessel_id=vessel.id,
            claim_reference="MCRI-HM-2026-0001",
            incident_date=date(2026, 7, 10),
            notification_date=date(2026, 7, 11),
            incident_description="Turbocharger failure",
            currency="USD",
        )
        db.add(claim)
        db.commit()
        return {"claim_id": str(claim.id), "alpha_org": str(alpha.id)}


def login(slug: str, email: str) -> None:
    response = client.post(
        "/api/v1/auth/login",
        json={"organization_slug": slug, "email": email, "password": PASSWORD},
    )
    assert response.status_code == 200, response.text


def configure_storage(tmp_path: Path) -> None:
    document_service.settings.local_storage_path = str(tmp_path / "documents")
    document_service.settings.max_upload_mb = 2


def make_docx_bytes() -> bytes:
    document = WordDocument()
    document.add_heading("Chief Engineer Report", level=1)
    document.add_paragraph("MT ORION experienced abnormal turbocharger vibration at 10:30 UTC.")
    document.add_paragraph("Engine load was reduced and the turbocharger was isolated.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_upload_auto_enqueues_text_job_and_worker_extracts_docx(tmp_path: Path) -> None:
    ids = seed_claim()
    configure_storage(tmp_path)
    login("alpha", "alpha@example.com")

    upload = client.post(
        f"/api/v1/claims/{ids['claim_id']}/documents",
        files={
            "file": (
                "CE_Report.docx",
                make_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"document_type": "chief_engineer_report"},
    )
    assert upload.status_code == 201, upload.text
    document_id = UUID(upload.json()["id"])

    with TestingSessionLocal() as db:
        job = db.scalar(select(DocumentProcessingJob).where(DocumentProcessingJob.document_id == document_id))
        assert job is not None
        assert job.status == ProcessingJobStatus.PENDING

        claimed = claim_next_job(db, worker_id="pytest-worker")
        assert claimed is not None
        assert claimed.id == job.id
        process_job(db, job=claimed)

    with TestingSessionLocal() as db:
        document = db.get(Document, document_id)
        assert document.processing_status == DocumentProcessingStatus.PROCESSED
        extraction = db.scalar(select(DocumentTextExtraction).where(DocumentTextExtraction.document_id == document_id))
        assert extraction is not None
        assert extraction.extraction_method == "python-docx"
        assert extraction.char_count > 50
        assert extraction.requires_ocr is False
        segments = list(db.scalars(select(DocumentTextSegment).where(DocumentTextSegment.document_id == document_id)))
        assert len(segments) == 1
        assert "abnormal turbocharger vibration" in segments[0].text

    summary = client.get(
        f"/api/v1/claims/{ids['claim_id']}/documents/{document_id}/processing"
    )
    assert summary.status_code == 200
    payload = summary.json()
    assert payload["job"]["status"] == "completed"
    assert payload["text_extraction"]["requires_ocr"] is False
    assert payload["text_extraction"]["char_count"] > 50


def test_image_processing_records_ocr_requirement(tmp_path: Path) -> None:
    ids = seed_claim()
    configure_storage(tmp_path)
    login("alpha", "alpha@example.com")
    png = b"\x89PNG\r\n\x1a\n" + b"placeholder-image-content"
    upload = client.post(
        f"/api/v1/claims/{ids['claim_id']}/documents",
        files={"file": ("damage.png", png, "image/png")},
    )
    assert upload.status_code == 201
    document_id = UUID(upload.json()["id"])

    with TestingSessionLocal() as db:
        claimed = claim_next_job(db, worker_id="pytest-worker")
        assert claimed is not None
        process_job(db, job=claimed)
        extraction = db.scalar(select(DocumentTextExtraction).where(DocumentTextExtraction.document_id == document_id))
        assert extraction is not None
        assert extraction.requires_ocr is True
        assert extraction.char_count == 0
        assert extraction.warnings


def test_processing_summary_is_tenant_protected(tmp_path: Path) -> None:
    ids = seed_claim()
    configure_storage(tmp_path)
    login("alpha", "alpha@example.com")
    upload = client.post(
        f"/api/v1/claims/{ids['claim_id']}/documents",
        files={"file": ("CE_Report.docx", make_docx_bytes(), "application/octet-stream")},
    )
    document_id = upload.json()["id"]

    client.cookies.clear()
    login("beta", "beta@example.com")
    response = client.get(
        f"/api/v1/claims/{ids['claim_id']}/documents/{document_id}/processing"
    )
    assert response.status_code == 404


def test_ai_gateway_is_provider_neutral_and_disabled_by_default() -> None:
    provider = get_ai_provider()
    assert provider.name == "disabled"
    request = AIRequest(
        task="classify_document",
        system_instructions="Classify only from supplied evidence.",
        input_text="Chief Engineer Report",
    )
    try:
        provider.generate(request)
        raise AssertionError("disabled provider should not generate")
    except AIProviderUnavailable:
        pass
