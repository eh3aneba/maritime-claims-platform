from datetime import date
from io import BytesIO
from pathlib import Path
from uuid import UUID

from docx import Document as WordDocument
from sqlalchemy import select

from app.modules.audit.models import AuditLog
from app.modules.claims.models import ClaimStatus
from app.modules.documents import service as document_service
from app.modules.processing import service as processing_service
from app.modules.processing.service import claim_next_job, process_job
from app.modules.rules.models import ClaimDocumentRequirement, RequirementStatus
from app.modules.tasks.models import ClaimTask, DocumentRequestBatch, TaskStatus
from tests.db_harness import TestingSessionLocal, client, reset_database
from tests.test_claims_api import create_orion_claim, login


def setup_function() -> None:
    reset_database()


def _set_status(claim_id: str, status: ClaimStatus) -> None:
    with TestingSessionLocal() as db:
        claim = db.get(__import__('app.modules.claims.models', fromlist=['Claim']).Claim, UUID(claim_id))
        claim.status = status
        db.commit()


def _evaluate(claim_id: str) -> dict:
    response = client.post(f"/api/v1/claims/{claim_id}/rules/evaluate")
    assert response.status_code == 200, response.text
    return response.json()["summary"]


def _configure_storage(tmp_path: Path) -> None:
    storage_path = str(tmp_path / "documents")
    document_service.settings.local_storage_path = storage_path
    document_service.settings.max_upload_mb = 1
    processing_service.settings.local_storage_path = storage_path


def _chief_engineer_docx() -> bytes:
    document = WordDocument()
    document.add_heading("Chief Engineer Report", level=1)
    document.add_paragraph("MT ORION experienced abnormal turbocharger vibration at 10:30 UTC.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _expected(item: dict) -> dict:
    return {
        "expected_state_fingerprint": item["state_fingerprint"],
        "expected_state_version": item["state_version"],
    }


def _approve_and_mark_sent(claim_id: str, batch_id: str) -> dict:
    items = client.get(f"/api/v1/claims/{claim_id}/correspondence").json()["items"]
    item = next(row for row in items if row["request_batch_id"] == batch_id)
    submitted = client.post(
        f"/api/v1/claims/{claim_id}/correspondence/{item['id']}/submit",
        json=_expected(item),
    )
    assert submitted.status_code == 200, submitted.text
    submitted_item = submitted.json()
    approved = client.post(
        f"/api/v1/claims/{claim_id}/correspondence/{item['id']}/approve",
        json={
            "note": "Reviewed against the claim file and approved for dispatch.",
            "confirm_re_review": False,
            **_expected(submitted_item),
        },
    )
    assert approved.status_code == 200, approved.text
    approved_item = approved.json()
    sent = client.post(
        f"/api/v1/claims/{claim_id}/correspondence/{item['id']}/mark-sent",
        json={
            "confirm_sent": True,
            "channel": "email",
            "external_reference": "OUT-2026-001",
            "expected_review_hash": approved_item["latest_review"]["review_hash"],
            **_expected(approved_item),
        },
    )
    assert sent.status_code == 200, sent.text
    return sent.json()


def test_request_all_critical_creates_draft_and_tasks_then_human_marks_requested() -> None:
    result = create_orion_claim()
    claim_id = result["claim"]["id"]
    _set_status(claim_id, ClaimStatus.TRIAGE)
    summary = _evaluate(claim_id)
    assert summary["readiness"]["critical_missing_count"] == 3

    response = client.post(
        f"/api/v1/claims/{claim_id}/document-requests",
        json={"all_critical": True, "due_date": "2026-08-12", "recipient_label": "Owners"},
    )
    assert response.status_code == 201, response.text
    payload = response.json()
    assert payload["batch"]["status"] == "draft"
    assert payload["batch"]["recipient_label"] == "Owners"
    assert "Chief Engineer Report" in payload["batch"]["draft_body"]
    assert len(payload["tasks"]) == 3
    assert all(task["status"] == "open" for task in payload["tasks"])
    assert all(task["source"] == "rule" for task in payload["tasks"])
    assert all(task["priority"] == "critical" for task in payload["tasks"])

    rules = client.get(f"/api/v1/claims/{claim_id}/rules").json()
    assert {r["status"] for r in rules["requirements"]} == {"missing"}  # Drafting is not the same as sending.
    marked = _approve_and_mark_sent(claim_id, payload["batch"]["id"])
    assert marked["status"] == "sent_externally"
    rules = client.get(f"/api/v1/claims/{claim_id}/rules").json()
    assert {r["status"] for r in rules["requirements"]} == {"requested"}
    assert rules["readiness"]["critical_missing_count"] == 3  # Requested is not satisfied.


def test_request_selected_requirement_does_not_duplicate_open_task() -> None:
    result = create_orion_claim(); claim_id = result["claim"]["id"]
    _set_status(claim_id, ClaimStatus.TRIAGE)
    summary = _evaluate(claim_id)
    ce = next(r for r in summary["requirements"] if r["document_type"] == "chief_engineer_report")
    for _ in range(2):
        response = client.post(f"/api/v1/claims/{claim_id}/document-requests", json={"requirement_ids": [ce["id"]]})
        assert response.status_code == 201, response.text
    tasks = client.get(f"/api/v1/claims/{claim_id}/tasks").json()
    assert tasks["total"] == 1
    with TestingSessionLocal() as db:
        batches = list(db.scalars(select(DocumentRequestBatch).where(DocumentRequestBatch.claim_id == UUID(claim_id))))
        assert len(batches) == 2  # Each draft is preserved for audit; the follow-up task is reused.


def test_document_upload_waits_for_processing_then_auto_completes_and_reopens(tmp_path: Path) -> None:
    result = create_orion_claim(); claim_id = result["claim"]["id"]
    _set_status(claim_id, ClaimStatus.TRIAGE); _configure_storage(tmp_path)
    summary = _evaluate(claim_id)
    ce = next(r for r in summary["requirements"] if r["document_type"] == "chief_engineer_report")
    request = client.post(f"/api/v1/claims/{claim_id}/document-requests", json={"requirement_ids": [ce["id"]], "due_date": "2026-08-12"})
    task_id = request.json()["tasks"][0]["id"]
    batch_id = request.json()["batch"]["id"]
    _approve_and_mark_sent(claim_id, batch_id)

    upload = client.post(
        f"/api/v1/claims/{claim_id}/documents",
        files={
            "file": (
                "CE_Report.docx",
                _chief_engineer_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"document_type": "chief_engineer_report", "confidentiality_level": "confidential"},
    )
    assert upload.status_code == 201, upload.text
    document_id = upload.json()["id"]

    pending_rules = client.get(f"/api/v1/claims/{claim_id}/rules").json()
    ce_pending = next(r for r in pending_rules["requirements"] if r["document_type"] == "chief_engineer_report")
    assert ce_pending["status"] == "received"
    assert ce_pending["satisfaction_basis"] == "document_processing_pending"
    assert pending_rules["readiness"]["critical_missing_count"] == 3
    pending_task = next(t for t in client.get(f"/api/v1/claims/{claim_id}/tasks").json()["items"] if t["id"] == task_id)
    assert pending_task["status"] == "open"

    with TestingSessionLocal() as db:
        claimed = claim_next_job(db, worker_id="phase-13-4a-test")
        assert claimed is not None
        process_job(db, job=claimed)

    processed_rules = client.get(f"/api/v1/claims/{claim_id}/rules").json()
    ce_processed = next(r for r in processed_rules["requirements"] if r["document_type"] == "chief_engineer_report")
    assert ce_processed["status"] == "under_review"
    assert ce_processed["satisfaction_basis"] in {"direct_document", "direct_document_legacy_unscanned"}
    assert processed_rules["readiness"]["critical_missing_count"] == 2
    completed_task = next(t for t in client.get(f"/api/v1/claims/{claim_id}/tasks").json()["items"] if t["id"] == task_id)
    assert completed_task["status"] == "completed"
    assert "available for human review" in completed_task["completion_reason"].lower()

    deleted = client.delete(f"/api/v1/claims/{claim_id}/documents/{document_id}")
    assert deleted.status_code == 204, deleted.text
    deleted_rules = client.get(f"/api/v1/claims/{claim_id}/rules").json()
    ce_deleted = next(r for r in deleted_rules["requirements"] if r["document_type"] == "chief_engineer_report")
    assert ce_deleted["status"] == "missing"
    reopened_task = next(t for t in client.get(f"/api/v1/claims/{claim_id}/tasks").json()["items"] if t["id"] == task_id)
    assert reopened_task["status"] == "open"
    assert reopened_task["completion_reason"] is None


def test_final_processing_failure_never_satisfies_or_closes_request(tmp_path: Path, monkeypatch) -> None:
    result = create_orion_claim(); claim_id = result["claim"]["id"]
    _set_status(claim_id, ClaimStatus.TRIAGE); _configure_storage(tmp_path)
    summary = _evaluate(claim_id)
    ce = next(r for r in summary["requirements"] if r["document_type"] == "chief_engineer_report")
    request = client.post(f"/api/v1/claims/{claim_id}/document-requests", json={"requirement_ids": [ce["id"]]})
    task_id = request.json()["tasks"][0]["id"]

    upload = client.post(
        f"/api/v1/claims/{claim_id}/documents",
        files={
            "file": (
                "CE_Report.docx",
                _chief_engineer_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"document_type": "chief_engineer_report"},
    )
    assert upload.status_code == 201, upload.text

    def fail_extraction(*args, **kwargs):
        del args, kwargs
        raise RuntimeError("synthetic extraction failure")

    monkeypatch.setattr(processing_service, "extract_document_text", fail_extraction)
    with TestingSessionLocal() as db:
        claimed = claim_next_job(db, worker_id="phase-13-4a-failure")
        assert claimed is not None
        claimed.max_attempts = claimed.attempt_count
        db.commit()
        process_job(db, job=claimed)

    failed_rules = client.get(f"/api/v1/claims/{claim_id}/rules").json()
    ce_failed = next(r for r in failed_rules["requirements"] if r["document_type"] == "chief_engineer_report")
    assert ce_failed["status"] == "received"
    assert ce_failed["satisfaction_basis"] == "document_processing_failed"
    assert failed_rules["readiness"]["critical_missing_count"] == 3
    task = next(t for t in client.get(f"/api/v1/claims/{claim_id}/tasks").json()["items"] if t["id"] == task_id)
    assert task["status"] == "open"


def test_manual_task_completion_is_audited_and_never_auto_reopened() -> None:
    result = create_orion_claim(); claim_id = result["claim"]["id"]
    _set_status(claim_id, ClaimStatus.TRIAGE)
    summary = _evaluate(claim_id)
    req = summary["requirements"][0]
    created = client.post(f"/api/v1/claims/{claim_id}/document-requests", json={"requirement_ids": [req["id"]]})
    task_id = created.json()["tasks"][0]["id"]
    completed = client.post(f"/api/v1/claims/{claim_id}/tasks/{task_id}/complete", json={"reason": "Owner confirmed this document is unavailable."})
    assert completed.status_code == 200
    assert completed.json()["status"] == "completed"
    _evaluate(claim_id)
    task = next(t for t in client.get(f"/api/v1/claims/{claim_id}/tasks").json()["items"] if t["id"] == task_id)
    assert task["status"] == "completed"
    assert task["completion_reason"] == "Owner confirmed this document is unavailable."
    with TestingSessionLocal() as db:
        audit = db.scalar(select(AuditLog).where(AuditLog.action == "COMPLETE_CLAIM_TASK", AuditLog.entity_id == UUID(task_id)))
        assert audit is not None


def test_document_requests_are_tenant_scoped() -> None:
    result = create_orion_claim(); claim_id = result["claim"]["id"]
    _set_status(claim_id, ClaimStatus.TRIAGE); _evaluate(claim_id)
    client.cookies.clear(); login("beta", "beta-handler@example.com")
    assert client.get(f"/api/v1/claims/{claim_id}/tasks").status_code == 404
    assert client.post(f"/api/v1/claims/{claim_id}/document-requests", json={"all_critical": True}).status_code == 404


def test_accepting_equivalent_evidence_auto_completes_open_requirement_task() -> None:
    from app.modules.claims.facts import ClaimFact
    from app.modules.claims.models import Claim

    result = create_orion_claim(); claim_id = result["claim"]["id"]
    _set_status(claim_id, ClaimStatus.INVESTIGATION)
    with TestingSessionLocal() as db:
        claim = db.get(Claim, UUID(claim_id))
        fact = ClaimFact(
            organization_id=claim.organization_id, claim_id=claim.id,
            field_path="maintenance.recommended_overhaul_interval", value={"value": 12000, "unit": "hours"},
            source_extraction_id=__import__('uuid').uuid4(), source_document_id=__import__('uuid').uuid4(), source_segment_id=None,
            approved_by_id=None, version=1,
        )
        db.add(fact); db.commit(); db.refresh(fact); fact_id=str(fact.id)
    summary = _evaluate(claim_id)
    req = next(r for r in summary["requirements"] if r["document_type"] == "maker_recommendation")
    candidate = next(row for row in req["equivalent_evidence_candidates"] if row["claim_fact_id"] == fact_id)
    request = client.post(f"/api/v1/claims/{claim_id}/document-requests", json={"requirement_ids": [req["id"]]})
    assert request.status_code == 201
    task_id = request.json()["tasks"][0]["id"]
    accepted = client.post(
        f"/api/v1/claims/{claim_id}/rules/requirements/{req['id']}/accept-equivalent",
        json={
            "claim_fact_id": fact_id,
            "claim_fact_version": candidate["claim_fact_version"],
            "expected_state_fingerprint": req["state_fingerprint"],
            "expected_state_version": req["state_version"],
            "note": "Reviewed approved maintenance interval as sufficient equivalent evidence.",
            "re_review": False,
        },
    )
    assert accepted.status_code == 200, accepted.text
    tasks = client.get(f"/api/v1/claims/{claim_id}/tasks").json()["items"]
    task = next(row for row in tasks if row["id"] == task_id)
    assert task["status"] == "completed"
    assert "equivalent evidence" in task["completion_reason"].lower()