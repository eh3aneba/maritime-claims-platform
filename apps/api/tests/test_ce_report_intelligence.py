from datetime import date
from io import BytesIO
from pathlib import Path
from uuid import UUID

from docx import Document as WordDocument
from sqlalchemy import select

from app.ai.gateway.base import AIRequest, AIResponse
from app.core.security import hash_password
from app.modules.claims.models import Claim
from app.modules.documents import service as document_service
from app.modules.documents.models import Document, DocumentProcessingStatus
from app.modules.intelligence.models import AIRun, AIRunStatus, AISemanticKind, AIReviewStatus, DocumentExtraction
from app.modules.intelligence.service import run_ce_report_intelligence
from app.modules.organizations.models import Organization
from app.modules.processing.models import DocumentProcessingJob, ProcessingJobStatus, ProcessingJobType
from app.modules.processing.service import claim_next_job, enqueue_processing_job, process_job
from app.modules.users.models import User, UserRole
from app.modules.vessels.models import Vessel
from tests.db_harness import TestingSessionLocal, client, reset_database

PASSWORD = "Strong-AI-Test-2026"


class FakeCEProvider:
    name = "fake"
    _model = "fake-maritime-v1"

    def __init__(self, *, bad_quote: bool = False, classification: str = "chief_engineer_report") -> None:
        self.bad_quote = bad_quote
        self.classification = classification
        self.last_request: AIRequest | None = None

    def generate(self, request: AIRequest) -> AIResponse:
        self.last_request = request
        quote = "text that does not exist" if self.bad_quote else "abnormal turbocharger vibration at 10:30 UTC"
        payload = {
            "classification": {"document_type": self.classification, "confidence": 0.98},
            "identification": {
                "vessel_name": sv("MT ORION", "MT ORION"),
                "imo_number": sv(None, None, 0),
                "report_date": sv("2026-07-10", "10 July 2026"),
                "author_name": sv(None, None, 0),
                "author_rank": sv("Chief Engineer", "Chief Engineer Report"),
            },
            "incident": {
                "date": sv("2026-07-10", "10 July 2026"),
                "time": sv("10:30", quote),
                "timezone": sv("UTC", "10:30 UTC"),
                "location": sv(None, None, 0),
                "voyage_from": sv(None, None, 0),
                "voyage_to": sv(None, None, 0),
                "cargo_status": sv(None, None, 0),
                "first_observation": sv("abnormal turbocharger vibration", quote),
            },
            "equipment": {
                "equipment_type": sv("turbocharger", "turbocharger vibration"),
                "equipment_name": sv("Main Engine Turbocharger", "turbocharger vibration"),
                "maker": sv(None, None, 0),
                "model": sv(None, None, 0),
                "serial_number": sv(None, None, 0),
            },
            "symptoms": [sv("abnormal vibration", quote)],
            "immediate_actions": [sv("Engine load was reduced", "Engine load was reduced")],
            "operational_impact": {
                "engine_stopped": sb(None, None, 0),
                "load_reduced": sb(True, "Engine load was reduced"),
                "speed_reduced": sb(None, None, 0),
                "immobilized": sb(None, None, 0),
                "deviation": sb(None, None, 0),
                "towage": sb(None, None, 0),
            },
            "suspected_cause_opinions": [sv("bearing damage", "I suspect bearing damage caused the failure")],
            "recommendations": [],
        }
        if self.classification != "chief_engineer_report":
            payload["symptoms"] = []
            payload["immediate_actions"] = []
            payload["suspected_cause_opinions"] = []
        return AIResponse(
            provider=self.name,
            model=self._model,
            structured_output=payload,
            output_text="{}",
            usage={"input_tokens": 100, "output_tokens": 80, "total_tokens": 180},
            raw_response_id="resp_test_001",
        )


def sv(value, quote, confidence: float = 0.95) -> dict:
    return {"value": value, "confidence": confidence, "source": {"segment_index": 0 if value is not None else None, "quote": quote}}


def sb(value, quote, confidence: float = 0.95) -> dict:
    return {"value": value, "confidence": confidence, "source": {"segment_index": 0 if value is not None else None, "quote": quote}}


def setup_function() -> None:
    reset_database()


def configure_storage(tmp_path: Path) -> None:
    document_service.settings.local_storage_path = str(tmp_path / "documents")
    document_service.settings.max_upload_mb = 2


def make_ce_docx() -> bytes:
    document = WordDocument()
    document.add_heading("Chief Engineer Report", level=1)
    document.add_paragraph("Report date: 10 July 2026")
    document.add_paragraph("MT ORION experienced abnormal turbocharger vibration at 10:30 UTC.")
    document.add_paragraph("Engine load was reduced and the turbocharger was isolated.")
    document.add_paragraph("I suspect bearing damage caused the failure.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def seed_claim() -> dict[str, str]:
    with TestingSessionLocal() as db:
        alpha = Organization(name="Alpha Marine", slug="alpha")
        beta = Organization(name="Beta Marine", slug="beta")
        db.add_all([alpha, beta])
        db.flush()
        alpha_user = User(
            organization_id=alpha.id,
            email="alpha@example.com",
            full_name="Alpha Handler",
            password_hash=hash_password(PASSWORD),
            role=UserRole.CLAIMS_HANDLER,
            is_active=True,
        )
        beta_user = User(
            organization_id=beta.id,
            email="beta@example.com",
            full_name="Beta Handler",
            password_hash=hash_password(PASSWORD),
            role=UserRole.CLAIMS_HANDLER,
            is_active=True,
        )
        vessel = Vessel(organization_id=alpha.id, name="MT ORION", imo_number="7000201")
        db.add_all([alpha_user, beta_user, vessel])
        db.flush()
        claim = Claim(
            organization_id=alpha.id,
            vessel_id=vessel.id,
            claim_reference="MCRI-HM-2026-0001",
            incident_date=date(2026, 7, 10),
            notification_date=date(2026, 7, 11),
            incident_description="Turbocharger failure",
            currency="USD",
        )
        db.add(claim)
        db.commit()
        return {"claim_id": str(claim.id), "org_id": str(alpha.id)}


def login(slug: str, email: str) -> None:
    response = client.post(
        "/api/v1/auth/login",
        json={"organization_slug": slug, "email": email, "password": PASSWORD},
    )
    assert response.status_code == 200, response.text


def upload_and_extract_text(tmp_path: Path, ids: dict[str, str]) -> UUID:
    configure_storage(tmp_path)
    login("alpha", "alpha@example.com")
    upload = client.post(
        f"/api/v1/claims/{ids['claim_id']}/documents",
        files={
            "file": (
                "CE_Report.docx",
                make_ce_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"document_type": "chief_engineer_report"},
    )
    assert upload.status_code == 201, upload.text
    document_id = UUID(upload.json()["id"])
    with TestingSessionLocal() as db:
        job = claim_next_job(db, worker_id="text-worker")
        assert job is not None and job.job_type == ProcessingJobType.EXTRACT_TEXT
        process_job(db, job=job)
    return document_id


def test_ce_report_ai_run_persists_source_linked_fact_and_opinion(tmp_path: Path) -> None:
    ids = seed_claim()
    document_id = upload_and_extract_text(tmp_path, ids)
    provider = FakeCEProvider()

    with TestingSessionLocal() as db:
        document = db.get(Document, document_id)
        run = run_ce_report_intelligence(db, document=document, requested_by_id=None, provider=provider)
        assert run.status == AIRunStatus.COMPLETED
        assert run.document_type_candidate == "chief_engineer_report"
        assert run.prompt_version == "1.0"
        assert run.schema_version == "1.0"
        assert run.raw_response_id == "resp_test_001"
        assert run.usage["total_tokens"] == 180

        rows = list(db.scalars(select(DocumentExtraction).where(DocumentExtraction.ai_run_id == run.id)))
        by_path = {row.field_path: row for row in rows}
        assert by_path["incident.time"].normalized_value == "10:30"
        assert by_path["incident.time"].semantic_kind == AISemanticKind.FACT
        assert by_path["incident.time"].source_verified is True
        assert by_path["suspected_cause_opinions[0]"].semantic_kind == AISemanticKind.OPINION
        assert by_path["suspected_cause_opinions[0]"].human_status == AIReviewStatus.PENDING
        assert by_path["operational_impact.load_reduced"].raw_value is True

        # AI classification/extraction must not silently overwrite the human/system document record.
        assert document.document_type == "chief_engineer_report"
        assert document.processing_status == DocumentProcessingStatus.PROCESSED

    assert provider.last_request is not None
    assert provider.last_request.output_schema is not None
    assert provider.last_request.schema_name == "chief_engineer_report_v1"
    assert "SEGMENT 0" in provider.last_request.input_text


def test_unverifiable_source_quote_is_flagged_not_silently_trusted(tmp_path: Path) -> None:
    ids = seed_claim()
    document_id = upload_and_extract_text(tmp_path, ids)
    with TestingSessionLocal() as db:
        document = db.get(Document, document_id)
        run = run_ce_report_intelligence(db, document=document, requested_by_id=None, provider=FakeCEProvider(bad_quote=True))
        row = db.scalar(
            select(DocumentExtraction).where(
                DocumentExtraction.ai_run_id == run.id,
                DocumentExtraction.field_path == "incident.time",
            )
        )
        assert row is not None
        assert row.source_verified is False
        assert row.validation_warnings
        assert any("incident.time" in warning for warning in (run.warnings or []))


def test_non_ce_classification_keeps_raw_run_but_does_not_persist_claim_facts(tmp_path: Path) -> None:
    ids = seed_claim()
    document_id = upload_and_extract_text(tmp_path, ids)
    with TestingSessionLocal() as db:
        document = db.get(Document, document_id)
        run = run_ce_report_intelligence(db, document=document, requested_by_id=None, provider=FakeCEProvider(classification="other"))
        assert run.status == AIRunStatus.COMPLETED
        assert run.document_type_candidate == "other"
        rows = list(db.scalars(select(DocumentExtraction).where(DocumentExtraction.ai_run_id == run.id)))
        assert rows == []
        assert run.raw_output is not None


def test_disabled_provider_blocks_api_enqueue(tmp_path: Path) -> None:
    ids = seed_claim()
    document_id = upload_and_extract_text(tmp_path, ids)
    response = client.post(f"/api/v1/claims/{ids['claim_id']}/documents/{document_id}/intelligence/ce-report")
    assert response.status_code == 409
    assert "disabled" in response.json()["detail"].lower()


def test_intelligence_summary_is_tenant_protected(tmp_path: Path) -> None:
    ids = seed_claim()
    document_id = upload_and_extract_text(tmp_path, ids)
    client.cookies.clear()
    login("beta", "beta@example.com")
    response = client.get(f"/api/v1/claims/{ids['claim_id']}/documents/{document_id}/intelligence")
    assert response.status_code == 404


def test_ai_job_dispatch_uses_background_queue_and_does_not_change_document_processing_status(tmp_path: Path, monkeypatch) -> None:
    ids = seed_claim()
    document_id = upload_and_extract_text(tmp_path, ids)
    provider = FakeCEProvider()
    monkeypatch.setattr("app.modules.intelligence.service.get_ai_provider", lambda: provider)

    with TestingSessionLocal() as db:
        document = db.get(Document, document_id)
        job = enqueue_processing_job(
            db,
            document=document,
            requested_by_id=None,
            job_type=ProcessingJobType.AI_EXTRACT_CE_REPORT,
        )
        db.commit()
        job_id = job.id

    with TestingSessionLocal() as db:
        claimed = claim_next_job(db, worker_id="ai-worker")
        assert claimed is not None
        assert claimed.id == job_id
        process_job(db, job=claimed)

    with TestingSessionLocal() as db:
        job = db.get(DocumentProcessingJob, job_id)
        document = db.get(Document, document_id)
        run = db.scalar(select(AIRun).where(AIRun.document_id == document_id))
        assert job.status == ProcessingJobStatus.COMPLETED
        assert run is not None and run.status == AIRunStatus.COMPLETED
        assert document.processing_status == DocumentProcessingStatus.PROCESSED


def test_openai_adapter_uses_responses_structured_output_shape() -> None:
    from app.ai.gateway.openai_provider import OpenAIProvider
    from app.ai.schemas.ce_report import ChiefEngineerReportExtraction

    captured = {}

    class Usage:
        input_tokens = 12
        output_tokens = 8
        total_tokens = 20

    class Response:
        id = "resp_openai_stub"
        output_text = '{"classification":{"document_type":"unknown","confidence":0.5},"identification":{"vessel_name":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"imo_number":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"report_date":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"author_name":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"author_rank":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}}},"incident":{"date":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"time":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"timezone":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"location":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"voyage_from":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"voyage_to":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"cargo_status":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"first_observation":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}}},"equipment":{"equipment_type":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"equipment_name":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"maker":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"model":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"serial_number":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}}},"symptoms":[],"immediate_actions":[],"operational_impact":{"engine_stopped":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"load_reduced":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"speed_reduced":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"immobilized":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"deviation":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}},"towage":{"value":null,"confidence":0,"source":{"segment_index":null,"quote":null}}},"suspected_cause_opinions":[],"recommendations":[]}'
        usage = Usage()

    class Responses:
        def create(self, **kwargs):
            captured.update(kwargs)
            return Response()

    class Client:
        responses = Responses()

    provider = OpenAIProvider.__new__(OpenAIProvider)
    provider._client = Client()
    provider._model = "configured-model"
    schema = ChiefEngineerReportExtraction.model_json_schema()
    response = provider.generate(
        AIRequest(
            task="chief_engineer_report_extract",
            system_instructions="extract only evidence",
            input_text="[SEGMENT 0] test",
            schema_name="chief_engineer_report_v1",
            output_schema=schema,
        )
    )
    assert captured["model"] == "configured-model"
    assert captured["text"]["format"]["type"] == "json_schema"
    assert captured["text"]["format"]["strict"] is True
    assert captured["text"]["format"]["schema"] == schema
    assert response.raw_response_id == "resp_openai_stub"
    assert response.usage["total_tokens"] == 20


def test_ce_schema_requires_all_top_level_fields_for_strict_output() -> None:
    from app.ai.schemas.ce_report import ChiefEngineerReportExtraction

    schema = ChiefEngineerReportExtraction.model_json_schema()
    expected = {
        "classification",
        "identification",
        "incident",
        "equipment",
        "symptoms",
        "immediate_actions",
        "operational_impact",
        "suspected_cause_opinions",
        "recommendations",
    }
    assert set(schema["required"]) == expected
    assert schema["additionalProperties"] is False


def test_restricted_document_external_ai_requires_separate_opt_in(tmp_path: Path, monkeypatch) -> None:
    ids = seed_claim()
    configure_storage(tmp_path)
    login("alpha", "alpha@example.com")
    upload = client.post(
        f"/api/v1/claims/{ids['claim_id']}/documents",
        files={
            "file": (
                "Restricted_CE.docx",
                make_ce_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"document_type": "chief_engineer_report", "confidentiality_level": "restricted"},
    )
    assert upload.status_code == 201
    document_id = UUID(upload.json()["id"])
    with TestingSessionLocal() as db:
        job = claim_next_job(db, worker_id="text-worker")
        assert job is not None
        process_job(db, job=job)

    class ExternalProvider:
        name = "openai"

    monkeypatch.setattr("app.modules.intelligence.router.get_ai_provider", lambda: ExternalProvider())
    monkeypatch.setattr("app.modules.intelligence.router.settings.allow_external_ai_restricted", False)
    response = client.post(
        f"/api/v1/claims/{ids['claim_id']}/documents/{document_id}/intelligence/ce-report"
    )
    assert response.status_code == 409
    assert "restricted" in response.json()["detail"].lower()
