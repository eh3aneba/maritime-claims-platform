from io import BytesIO
from pathlib import Path

from app.core.security import hash_password
from app.modules.audit.models import AuditLog
from app.modules.claims.facts import ClaimFact
from app.modules.organizations.models import Organization
from app.modules.users.models import User, UserRole
from docx import Document as WordDocument
from sqlalchemy import func, select

from app.modules.claims.models import Claim
from app.modules.documents.malware import MalwareScanResult, MalwareScanVerdict
from app.modules.documents.models import Document, DocumentMalwareScanStatus
from app.modules.intake import service as intake_service
from app.modules.intake.models import ClaimIntakeDraft, ClaimIntakeStatus
from app.modules.vessels.models import Vessel
from tests.db_harness import TestingSessionLocal, client, reset_database

TEST_LOGIN_VALUE = "test-only-value"


def setup_function() -> None:
    reset_database()


def seed() -> dict[str, str]:
    with TestingSessionLocal() as db:
        alpha = Organization(name="Alpha Marine", slug="alpha")
        beta = Organization(name="Beta Marine", slug="beta")
        db.add_all([alpha, beta])
        db.flush()
        alpha_user = User(
            organization_id=alpha.id,
            email="alpha@example.com",
            full_name="Alpha Manager",
            password_hash=hash_password(TEST_LOGIN_VALUE),
            role=UserRole.CLAIMS_MANAGER,
            is_active=True,
        )
        beta_user = User(
            organization_id=beta.id,
            email="beta@example.com",
            full_name="Beta Manager",
            password_hash=hash_password(TEST_LOGIN_VALUE),
            role=UserRole.CLAIMS_MANAGER,
            is_active=True,
        )
        alpha_vessel = Vessel(organization_id=alpha.id, name="MT ORION", imo_number="7654321")
        beta_vessel = Vessel(organization_id=beta.id, name="MT BETA", imo_number="7654322")
        db.add_all([alpha_user, beta_user, alpha_vessel, beta_vessel])
        db.commit()
        return {
            "alpha_org": str(alpha.id),
            "alpha_vessel": str(alpha_vessel.id),
            "beta_vessel": str(beta_vessel.id),
        }


def login(slug: str, email: str) -> None:
    response = client.post(
        "/api/v1/auth/login",
        json={"organization_slug": slug, "email": email, "password": TEST_LOGIN_VALUE},
    )
    assert response.status_code == 200, response.text


def configure_intake(tmp_path: Path, monkeypatch, *, infected: bool = False) -> None:
    monkeypatch.setattr(intake_service.settings, "local_storage_path", str(tmp_path / "documents"))
    monkeypatch.setattr(intake_service.settings, "max_upload_mb", 2)
    monkeypatch.setattr(intake_service.settings, "malware_scan_enabled", True)
    monkeypatch.setattr(intake_service.settings, "ocr_enabled", False)
    verdict = MalwareScanVerdict.INFECTED if infected else MalwareScanVerdict.CLEAN
    monkeypatch.setattr(
        intake_service,
        "scan_file",
        lambda *args, **kwargs: MalwareScanResult(
            verdict=verdict,
            threat_name="Eicar-Test-Signature" if infected else None,
        ),
    )


def make_fnol() -> bytes:
    document = WordDocument()
    document.add_heading("Claim Notification", level=1)
    document.add_paragraph("Vessel Name: MT ORION")
    document.add_paragraph("IMO Number: 7654321")
    document.add_paragraph("Incident Date: 2026-08-10")
    document.add_paragraph("Notification Date: 2026-08-11")
    document.add_paragraph("Claim Reference: CLUB-2026-42")
    document.add_paragraph(
        "Incident Description: Main engine turbocharger developed abnormal vibration during voyage."
    )
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def upload_and_process(tmp_path: Path, monkeypatch) -> dict:
    configure_intake(tmp_path, monkeypatch)
    response = client.post(
        "/api/v1/claim-intake/drafts",
        files={
            "file": (
                "claim_notification.docx",
                make_fnol(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 202, response.text
    assert response.json()["status"] == "processing"

    with TestingSessionLocal() as db:
        assert db.scalar(select(func.count(Claim.id))) == 0
        job = intake_service.claim_next_intake_job(db, worker_id="pytest-intake")
        assert job is not None
        intake_service.process_intake_job(db, job=job)

    reviewed = client.get(f"/api/v1/claim-intake/drafts/{response.json()['id']}")
    assert reviewed.status_code == 200
    assert reviewed.json()["status"] == "pending_review"
    return reviewed.json()


def approval_payload(vessel_id: str, extracted: dict) -> dict:
    return {
        "claim": {
            "vessel_id": vessel_id,
            "incident_date": extracted["incident_date"],
            "notification_date": extracted["notification_date"],
            "incident_description": extracted["incident_description"],
            "claim_type": "hull_machinery",
            "claim_subtype": "machinery_damage",
            "priority": "medium",
            "external_reference": extracted["external_reference"],
            "currency": "USD",
        },
        "document_type": "claim_notification",
        "review_note": "Reviewed against the uploaded FNOL and approved by a claims manager.",
    }


def test_fnol_candidates_require_review_before_exactly_one_claim_is_created(
    tmp_path: Path,
    monkeypatch,
) -> None:
    ids = seed()
    login("alpha", "alpha@example.com")
    draft = upload_and_process(tmp_path, monkeypatch)
    assert draft["classification_candidate"] == "claim_notification"
    assert draft["classification_confidence"] >= 90
    assert draft["extracted_fields"]["vessel_name"] == "MT ORION"
    assert draft["extracted_fields"]["imo_number"] == "7654321"

    payload = approval_payload(ids["alpha_vessel"], draft["extracted_fields"])
    first = client.post(f"/api/v1/claim-intake/drafts/{draft['id']}/approve", json=payload)
    assert first.status_code == 200, first.text
    second = client.post(f"/api/v1/claim-intake/drafts/{draft['id']}/approve", json=payload)
    assert second.status_code == 200, second.text
    assert first.json()["claim"]["id"] == second.json()["claim"]["id"]
    assert first.json()["draft"]["status"] == "approved"

    with TestingSessionLocal() as db:
        assert db.scalar(select(func.count(Claim.id))) == 1
        assert db.scalar(select(func.count(Document.id))) == 1
        assert db.scalar(select(func.count(ClaimFact.id))) == 0
        source = db.scalar(select(Document))
        assert source is not None
        assert source.malware_scan_status == DocumentMalwareScanStatus.CLEAN
        assert source.document_type == "claim_notification"
        actions = set(db.scalars(select(AuditLog.action)))
        assert "CREATE_CLAIM_INTAKE_DRAFT" in actions
        assert "EXTRACT_CLAIM_INTAKE_CANDIDATES" in actions
        assert "APPROVE_CLAIM_INTAKE_DRAFT" in actions
        assert "CREATE_CLAIM_FROM_INTAKE" in actions


def test_intake_draft_is_tenant_scoped_and_cross_tenant_approval_is_hidden(
    tmp_path: Path,
    monkeypatch,
) -> None:
    ids = seed()
    login("alpha", "alpha@example.com")
    draft = upload_and_process(tmp_path, monkeypatch)
    client.cookies.clear()
    login("beta", "beta@example.com")

    assert client.get(f"/api/v1/claim-intake/drafts/{draft['id']}").status_code == 404
    forbidden = client.post(
        f"/api/v1/claim-intake/drafts/{draft['id']}/approve",
        json=approval_payload(ids["beta_vessel"], draft["extracted_fields"]),
    )
    assert forbidden.status_code == 404


def test_infected_intake_is_quarantined_and_never_creates_claim(
    tmp_path: Path, monkeypatch
) -> None:
    seed()
    login("alpha", "alpha@example.com")
    configure_intake(tmp_path, monkeypatch, infected=True)
    response = client.post(
        "/api/v1/claim-intake/drafts",
        files={"file": ("infected.docx", make_fnol(), "application/octet-stream")},
    )
    assert response.status_code == 422
    with TestingSessionLocal() as db:
        draft = db.scalar(select(ClaimIntakeDraft))
        assert draft is not None
        assert draft.status == ClaimIntakeStatus.INFECTED
        assert draft.malware_scan_status == DocumentMalwareScanStatus.INFECTED_QUARANTINED
        assert db.scalar(select(func.count(Claim.id))) == 0
        assert db.scalar(select(func.count(Document.id))) == 0


def test_rejected_intake_never_creates_claim(tmp_path: Path, monkeypatch) -> None:
    seed()
    login("alpha", "alpha@example.com")
    draft = upload_and_process(tmp_path, monkeypatch)
    response = client.post(
        f"/api/v1/claim-intake/drafts/{draft['id']}/reject",
        json={"reason": "The uploaded notification belongs to a different casualty record."},
    )
    assert response.status_code == 200
    assert response.json()["status"] == "rejected"
    with TestingSessionLocal() as db:
        assert db.scalar(select(func.count(Claim.id))) == 0
        assert db.scalar(select(func.count(Document.id))) == 0


def test_candidate_extraction_is_deterministic_and_non_authoritative() -> None:
    fields, evidence = intake_service.extract_intake_candidates(
        "Claim Notification\nVessel Name: MT ORION\nIMO: 7654321\nIncident Date: 10 August 2026\n"
        "Notification Date: 2026-08-11\nIncident Description: Turbocharger vibration required load reduction."
    )
    assert fields["incident_date"] == "2026-08-10"
    assert fields["notification_date"] == "2026-08-11"
    assert evidence["priority"]["confidence"] == 0
    assert "human review" in evidence["priority"]["note"]
